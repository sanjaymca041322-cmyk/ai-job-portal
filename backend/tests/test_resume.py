from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.core.config import settings
from app.db.session import get_db
from app.main import app
from app.modules.candidates.parser import ResumeParser
from app.modules.candidates.resume_service import extract_text, parse_resume_text


def candidate_payload() -> dict[str, str]:
    return {
        "first_name": "Ada",
        "last_name": "Lovelace",
        "email": "ada@example.com",
    }


@pytest.fixture
def client(tmp_path: Path):
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from sqlalchemy.pool import StaticPool
    from app.db.base import Base

    engine = create_engine("sqlite:///:memory:", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    Base.metadata.create_all(engine)
    session = sessionmaker(bind=engine)()
    def override_get_db():
        yield session

    app.dependency_overrides[get_db] = override_get_db
    old_directory = settings.upload_directory
    settings.upload_directory = str(tmp_path)
    yield TestClient(app), session
    settings.upload_directory = old_directory
    app.dependency_overrides.clear()
    session.close()
    Base.metadata.drop_all(engine)


def create_candidate(client: TestClient) -> str:
    return client.post("/candidates", json=candidate_payload()).json()["id"]


def test_pdf_upload(client):
    test_client, _ = client
    candidate_id = create_candidate(test_client)
    response = test_client.post(
        f"/candidates/{candidate_id}/resume",
        files={"file": ("resume.pdf", b"%PDF-1.4 resume", "application/pdf")},
    )
    assert response.status_code == 201
    assert response.json()["resume_filename"] == "resume.pdf"


def test_docx_upload(client):
    test_client, _ = client
    candidate_id = create_candidate(test_client)
    response = test_client.post(
        f"/candidates/{candidate_id}/resume",
        files={"file": ("resume.docx", b"docx bytes", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 201


def test_upload_rejects_unsupported_empty_and_missing_candidate(client):
    test_client, _ = client
    candidate_id = create_candidate(test_client)
    assert test_client.post(f"/candidates/{candidate_id}/resume", files={"file": ("resume.txt", b"text", "text/plain")}).status_code == 400
    assert test_client.post(f"/candidates/{candidate_id}/resume", files={"file": ("resume.pdf", b"", "application/pdf")}).status_code == 400
    assert test_client.post(f"/candidates/{candidate_id}/resume", files={"file": ("resume.pdf", b"x", "application/pdf")}).status_code == 201


def test_resume_parse_and_profile_update(client):
    test_client, session = client
    candidate_id = create_candidate(test_client)
    document = Document()
    document.add_paragraph("Ada Lovelace")
    document.add_paragraph("ada.resume@example.com")
    document.add_paragraph("5 years experience")
    document.add_paragraph("Skills: Python, SQL")
    document.add_paragraph("Education: Mathematics")
    from io import BytesIO
    buffer = BytesIO()
    document.save(buffer)
    upload = test_client.post(f"/candidates/{candidate_id}/resume", files={"file": ("resume.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert upload.status_code == 201
    response = test_client.post(f"/candidates/{candidate_id}/resume/parse")
    assert response.status_code == 200
    assert response.json()["email"] == "ada.resume@example.com"
    assert response.json()["experience_years"] == 5
    candidate = session.query(__import__("app.modules.candidates.model", fromlist=["Candidate"]).Candidate).first()
    assert candidate.email == "ada.resume@example.com"


def test_parse_missing_and_corrupt_resume(client):
    test_client, _ = client
    candidate_id = create_candidate(test_client)
    assert test_client.post(f"/candidates/{candidate_id}/resume/parse").status_code == 400
    response = test_client.post(f"/candidates/{candidate_id}/resume", files={"file": ("resume.pdf", b"bad", "application/pdf")})
    assert response.status_code == 201
    assert test_client.post(f"/candidates/{candidate_id}/resume/parse").status_code == 422


def test_upload_missing_candidate(client):
    test_client, _ = client
    response = test_client.post(
        "/candidates/00000000-0000-0000-0000-000000000000/resume",
        files={"file": ("resume.pdf", b"resume", "application/pdf")},
    )
    assert response.status_code == 404


def test_extract_docx_text(tmp_path: Path):
    document = Document()
    document.add_paragraph("Ada Lovelace")
    document.add_paragraph("Skills: Python, SQL")
    path = tmp_path / "resume.docx"
    document.save(path)

    assert "Ada Lovelace" in extract_text(str(path))
    assert "Python, SQL" in extract_text(str(path))


def test_parse_resume_text_returns_structured_data():
    parsed = parse_resume_text(
        "Ada Lovelace\nada@example.com\n5 years experience\nSkills: Python, SQL"
    )

    assert parsed.name == "Ada Lovelace"
    assert parsed.email == "ada@example.com"
    assert parsed.experience_years == 5
    assert parsed.skills == ["Python", "SQL"]


def test_fallback_parser_works_without_api_key(monkeypatch):
    monkeypatch.setattr(settings, "llm_api_key", None)
    parser = ResumeParser()

    parsed = parser.parse("Ada Lovelace\nada@example.com")

    assert parsed.email == "ada@example.com"


def test_corrupt_docx_extraction_fails(tmp_path: Path):
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a valid docx")

    with pytest.raises(Exception):
        extract_text(str(path))